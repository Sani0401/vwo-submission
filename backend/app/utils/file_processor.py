"""
File Processing Utility for Financial Document Analyzer
Handles multiple file types: PDF, DOC, DOCX, XLS, XLSX
"""

import os
import io
from typing import Optional, Union
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileProcessor:
    """Utility class to process different types of financial documents"""
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.txt'}
    
    def is_supported_file(self, file_path: str) -> bool:
        """Check if the file type is supported"""
        _, ext = os.path.splitext(file_path.lower())
        return ext in self.supported_extensions
    
    def get_file_extension(self, file_path: str) -> str:
        """Get file extension"""
        _, ext = os.path.splitext(file_path.lower())
        return ext
    
    def process_file(self, file_path: str) -> str:
        """
        Process file based on its type and extract text content
        
        Args:
            file_path (str): Path to the file to process
            
        Returns:
            str: Extracted text content from the file
        """
        try:
            if not os.path.exists(file_path):
                return f"Error: File {file_path} not found"
            
            if not self.is_supported_file(file_path):
                return f"Error: Unsupported file type. Supported types: {', '.join(self.supported_extensions)}"
            
            file_ext = self.get_file_extension(file_path)
            
            if file_ext == '.pdf':
                return self._process_pdf(file_path)
            elif file_ext in ['.doc', '.docx']:
                return self._process_word_document(file_path)
            elif file_ext in ['.xls', '.xlsx']:
                return self._process_excel(file_path)
            elif file_ext == '.txt':
                return self._process_text_file(file_path)
            else:
                return f"Error: No processor available for {file_ext} files"
                
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return f"Error processing file: {str(e)}"
    
    def _process_pdf(self, file_path: str) -> str:
        """Process PDF files and extract text"""
        try:
            import PyPDF2
            
            text_content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n"
            
            if not text_content.strip():
                return "Error: No text content could be extracted from the PDF"
            
            return text_content.strip()
            
        except ImportError:
            return "Error: PyPDF2 library not installed. Please install it with: pip install PyPDF2"
        except Exception as e:
            return f"Error reading PDF file: {str(e)}"
    
    def _process_word_document(self, file_path: str) -> str:
        """Process Word documents (.doc, .docx) and extract text"""
        try:
            import docx
            
            if file_path.lower().endswith('.docx'):
                doc = docx.Document(file_path)
                text_content = ""
                
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_content += cell.text + " "
                        text_content += "\n"
                
                return text_content.strip()
            
            else:  # .doc files
                try:
                    import win32com.client
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(file_path)
                    text_content = doc.Content.Text
                    doc.Close()
                    word.Quit()
                    return text_content.strip()
                except ImportError:
                    return "Error: For .doc files, please install pywin32 library or convert to .docx format"
                    
        except ImportError:
            return "Error: python-docx library not installed. Please install it with: pip install python-docx"
        except Exception as e:
            return f"Error reading Word document: {str(e)}"
    
    def _process_excel(self, file_path: str) -> str:
        """Process Excel files (.xls, .xlsx) and extract text"""
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_content = ""
            
            for sheet_name in excel_file.sheet_names:
                text_content += f"\n--- Sheet: {sheet_name} ---\n"
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert DataFrame to text representation
                text_content += df.to_string(index=False) + "\n"
                
                # Also extract any text-like content
                for column in df.columns:
                    if df[column].dtype == 'object':  # Text columns
                        text_values = df[column].dropna().astype(str)
                        if len(text_values) > 0:
                            text_content += f"\n{column} values: " + " ".join(text_values.head(10)) + "\n"
            
            return text_content.strip()
            
        except ImportError:
            return "Error: pandas library not installed. Please install it with: pip install pandas openpyxl"
        except Exception as e:
            return f"Error reading Excel file: {str(e)}"
    
    def _process_text_file(self, file_path: str) -> str:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return content.strip()
        except UnicodeDecodeError:
            # Try with different encodings
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    return content.strip()
                except UnicodeDecodeError:
                    continue
            return "Error: Could not decode text file with any supported encoding"
        except Exception as e:
            return f"Error reading text file: {str(e)}"
    
    def get_file_info(self, file_path: str) -> dict:
        """Get basic information about the file"""
        try:
            if not os.path.exists(file_path):
                return {"error": "File not found"}
            
            stat = os.stat(file_path)
            file_ext = self.get_file_extension(file_path)
            
            return {
                "filename": os.path.basename(file_path),
                "extension": file_ext,
                "size_bytes": stat.st_size,
                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                "is_supported": self.is_supported_file(file_path),
                "modified_time": stat.st_mtime
            }
        except Exception as e:
            return {"error": str(e)}

# Create a global instance
file_processor = FileProcessor()

def process_financial_document(file_path: str) -> str:
    """
    Main function to process financial documents of any supported type
    
    Args:
        file_path (str): Path to the financial document
        
    Returns:
        str: Extracted text content ready for analysis
    """
    return file_processor.process_file(file_path)

def get_document_info(file_path: str) -> dict:
    """
    Get information about a document file
    
    Args:
        file_path (str): Path to the document
        
    Returns:
        dict: Document information
    """
    return file_processor.get_file_info(file_path)
