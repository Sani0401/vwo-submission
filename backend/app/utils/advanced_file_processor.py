"""
Advanced File Processing Utility for Financial Document Analyzer
Handles edge cases, large files, corrupted files, password-protected files, and OCR
"""

import os
import io
import hashlib
import tempfile
from typing import Optional, Union, Dict, Any, List, Tuple
import logging
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor
import time

# Enhanced file processing dependencies
try:
    import PyPDF2
    import pdfplumber
    from pdfplumber import PDF
    PYPDF2_AVAILABLE = True
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    PDFPLUMBER_AVAILABLE = False
    logging.warning("PyPDF2/pdfplumber not available. PDF processing will be limited.")

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    logging.warning("python-magic not available. File type detection will be limited.")

try:
    import python_docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX processing will be limited.")

try:
    import openpyxl
    import pandas as pd
    OPENPYXL_AVAILABLE = True
    PANDAS_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    PANDAS_AVAILABLE = False
    logging.warning("openpyxl/pandas not available. Excel processing will be limited.")

try:
    import pytesseract
    from PIL import Image
    import fitz  # PyMuPDF
    TESSERACT_AVAILABLE = True
    PYMUPDF_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    PYMUPDF_AVAILABLE = False
    logging.warning("Tesseract/PyMuPDF not available. OCR processing will be limited.")

try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logging.warning("pypdf not available. Advanced PDF processing will be limited.")

from app.models.schemas import LogLevel, LogCategory, LogAction

# Set up logging
logger = logging.getLogger(__name__)

class AdvancedFileProcessor:
    """Advanced file processor with edge case handling"""
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.txt', '.png', '.jpg', '.jpeg', '.tiff', '.bmp'}
        self.max_file_size = 100 * 1024 * 1024  # 100MB
        self.max_pages = 1000
        self.ocr_languages = ['eng', 'fra', 'deu', 'spa', 'ita', 'por', 'rus', 'chi_sim', 'jpn', 'kor']
        
        # Initialize logging service
        from app.core.logging import get_logging_service
        self.logging_service = get_logging_service()
    
    def validate_file(self, file_path: str, file_size: int = None) -> Dict[str, Any]:
        """Comprehensive file validation"""
        validation_result = {
            'is_valid': True,
            'errors': [],
            'warnings': [],
            'file_info': {},
            'processing_recommendations': []
        }
        
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                validation_result['is_valid'] = False
                validation_result['errors'].append("File does not exist")
                return validation_result
            
            # Get file size
            if file_size is None:
                file_size = os.path.getsize(file_path)
            
            validation_result['file_info']['size_bytes'] = file_size
            validation_result['file_info']['size_mb'] = file_size / (1024 * 1024)
            
            # Check file size
            if file_size > self.max_file_size:
                validation_result['is_valid'] = False
                validation_result['errors'].append(f"File too large: {file_size / (1024 * 1024):.2f}MB (max: {self.max_file_size / (1024 * 1024):.2f}MB)")
            elif file_size > 50 * 1024 * 1024:  # 50MB warning
                validation_result['warnings'].append(f"Large file: {file_size / (1024 * 1024):.2f}MB - processing may be slow")
            
            # Check file extension
            file_ext = Path(file_path).suffix.lower()
            validation_result['file_info']['extension'] = file_ext
            
            if file_ext not in self.supported_extensions:
                validation_result['is_valid'] = False
                validation_result['errors'].append(f"Unsupported file type: {file_ext}")
                return validation_result
            
            # Detect actual file type using magic numbers
            try:
                if MAGIC_AVAILABLE:
                    mime_type = magic.from_file(file_path, mime=True)
                    validation_result['file_info']['mime_type'] = mime_type
                else:
                    validation_result['file_info']['mime_type'] = 'unknown'
                
                # Check if extension matches actual file type
                expected_mime_types = {
                    '.pdf': 'application/pdf',
                    '.doc': 'application/msword',
                    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    '.xls': 'application/vnd.ms-excel',
                    '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    '.txt': 'text/plain',
                    '.png': 'image/png',
                    '.jpg': 'image/jpeg',
                    '.jpeg': 'image/jpeg',
                    '.tiff': 'image/tiff',
                    '.bmp': 'image/bmp'
                }
                
                if file_ext in expected_mime_types:
                    expected_mime = expected_mime_types[file_ext]
                    if mime_type != expected_mime:
                        validation_result['warnings'].append(f"File extension ({file_ext}) doesn't match actual file type ({mime_type})")
            except Exception as e:
                validation_result['warnings'].append(f"Could not detect MIME type: {str(e)}")
            
            # File-specific validations
            if file_ext == '.pdf':
                self._validate_pdf(file_path, validation_result)
            elif file_ext in ['.doc', '.docx']:
                self._validate_document(file_path, validation_result)
            elif file_ext in ['.xls', '.xlsx']:
                self._validate_spreadsheet(file_path, validation_result)
            elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                self._validate_image(file_path, validation_result)
            
            return validation_result
            
        except Exception as e:
            validation_result['is_valid'] = False
            validation_result['errors'].append(f"Validation error: {str(e)}")
            return validation_result
    
    def _validate_pdf(self, file_path: str, validation_result: Dict[str, Any]):
        """Validate PDF file"""
        try:
            if PYPDF2_AVAILABLE:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    # Check if PDF is encrypted
                    if pdf_reader.is_encrypted:
                        validation_result['is_valid'] = False
                        validation_result['errors'].append("PDF is password-protected")
                        validation_result['processing_recommendations'].append("Password required for processing")
                        return
                    
                    # Check number of pages
                    num_pages = len(pdf_reader.pages)
                    validation_result['file_info']['num_pages'] = num_pages
                    
                    if num_pages > self.max_pages:
                        validation_result['is_valid'] = False
                        validation_result['errors'].append(f"Too many pages: {num_pages} (max: {self.max_pages})")
                    elif num_pages > 100:
                        validation_result['warnings'].append(f"Large document: {num_pages} pages - processing may be slow")
                    
                    # Check for text content
                    has_text = False
                    for page_num in range(min(5, num_pages)):  # Check first 5 pages
                        page = pdf_reader.pages[page_num]
                        if page.extract_text().strip():
                            has_text = True
                            break
                    
                    if not has_text:
                        validation_result['warnings'].append("PDF appears to be image-based (scanned document)")
                        validation_result['processing_recommendations'].append("OCR processing recommended")
                    
        except Exception as e:
            validation_result['warnings'].append(f"PDF validation error: {str(e)}")
    
    def _validate_document(self, file_path: str, validation_result: Dict[str, Any]):
        """Validate Word document"""
        try:
            if DOCX_AVAILABLE and file_path.endswith('.docx'):
                doc = Document(file_path)
                validation_result['file_info']['num_paragraphs'] = len(doc.paragraphs)
                
                # Check for tables
                num_tables = len(doc.tables)
                validation_result['file_info']['num_tables'] = num_tables
                
                if num_tables > 0:
                    validation_result['processing_recommendations'].append("Document contains tables - structured extraction recommended")
        except Exception as e:
            validation_result['warnings'].append(f"Document validation error: {str(e)}")
    
    def _validate_spreadsheet(self, file_path: str, validation_result: Dict[str, Any]):
        """Validate Excel spreadsheet"""
        try:
            if OPENPYXL_AVAILABLE and file_path.endswith('.xlsx'):
                workbook = openpyxl.load_workbook(file_path, read_only=True)
                validation_result['file_info']['num_sheets'] = len(workbook.sheetnames)
                validation_result['file_info']['sheet_names'] = workbook.sheetnames
                
                if len(workbook.sheetnames) > 10:
                    validation_result['warnings'].append(f"Many worksheets: {len(workbook.sheetnames)} - processing may be slow")
        except Exception as e:
            validation_result['warnings'].append(f"Spreadsheet validation error: {str(e)}")
    
    def _validate_image(self, file_path: str, validation_result: Dict[str, Any]):
        """Validate image file"""
        try:
            if TESSERACT_AVAILABLE:
                with Image.open(file_path) as img:
                    validation_result['file_info']['image_size'] = img.size
                    validation_result['file_info']['image_mode'] = img.mode
                    
                    # Check image quality
                    if img.size[0] * img.size[1] < 10000:  # Very small image
                        validation_result['warnings'].append("Very small image - OCR quality may be poor")
                    
                    validation_result['processing_recommendations'].append("OCR processing required")
        except Exception as e:
            validation_result['warnings'].append(f"Image validation error: {str(e)}")
    
    async def process_file_advanced(
        self,
        file_path: str,
        user_id: str = None,
        options: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """Advanced file processing with edge case handling"""
        
        if options is None:
            options = {}
        
        processing_result = {
            'success': False,
            'content': '',
            'metadata': {},
            'processing_info': {},
            'errors': [],
            'warnings': []
        }
        
        try:
            # Validate file first
            validation = self.validate_file(file_path)
            processing_result['validation'] = validation
            
            if not validation['is_valid']:
                processing_result['errors'] = validation['errors']
                return processing_result
            
            processing_result['warnings'] = validation['warnings']
            processing_result['metadata'] = validation['file_info']
            
            # Log processing start
            self.logging_service.log_activity(
                level=LogLevel.INFO,
                category=LogCategory.DOCUMENT,
                action=LogAction.DOCUMENT_UPLOAD,
                message=f"Starting advanced file processing: {Path(file_path).name}",
                user_id=user_id,
                details={
                    'file_size': validation['file_info']['size_bytes'],
                    'file_type': validation['file_info']['extension']
                }
            )
            
            start_time = time.time()
            file_ext = validation['file_info']['extension']
            
            # Process based on file type
            if file_ext == '.pdf':
                content = self._process_pdf_advanced(file_path, options)
            elif file_ext in ['.doc', '.docx']:
                content = self._process_document_advanced(file_path, options)
            elif file_ext in ['.xls', '.xlsx']:
                content = self._process_spreadsheet_advanced(file_path, options)
            elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                content = self._process_image_advanced(file_path, options)
            elif file_ext == '.txt':
                content = self._process_text_advanced(file_path, options)
            else:
                processing_result['errors'].append(f"Unsupported file type: {file_ext}")
                return processing_result
            
            processing_time = time.time() - start_time
            processing_result['processing_info']['processing_time_seconds'] = processing_time
            processing_result['processing_info']['content_length'] = len(content)
            
            if content:
                processing_result['success'] = True
                processing_result['content'] = content
                
                # Log successful processing
                self.logging_service.log_activity(
                    level=LogLevel.INFO,
                    category=LogCategory.DOCUMENT,
                    action=LogAction.DOCUMENT_UPLOAD,
                    message=f"File processed successfully: {Path(file_path).name}",
                    user_id=user_id,
                    details={
                        'processing_time': processing_time,
                        'content_length': len(content)
                    }
                )
            else:
                processing_result['errors'].append("No content extracted from file")
            
            return processing_result
            
        except Exception as e:
            processing_result['errors'].append(f"Processing error: {str(e)}")
            logger.error(f"Error processing file {file_path}: {e}")
            
            # Log processing error
            self.logging_service.log_activity(
                level=LogLevel.ERROR,
                category=LogCategory.DOCUMENT,
                action=LogAction.DOCUMENT_UPLOAD,
                message=f"File processing failed: {Path(file_path).name}",
                user_id=user_id,
                details={'error': str(e)}
            )
            
            return processing_result
    
    async def _process_pdf_advanced(self, file_path: str, options: Dict[str, Any]) -> str:
        """Advanced PDF processing with multiple extraction methods"""
        content = ""
        
        try:
            # Try pdfplumber first (better for tables and layout)
            if PDFPLUMBER_AVAILABLE:
                try:
                    with PDF.open(file_path) as pdf:
                        for page in pdf.pages:
                            # Extract text
                            page_text = page.extract_text()
                            if page_text:
                                content += page_text + "\n"
                            
                            # Extract tables
                            tables = page.extract_tables()
                            for table in tables:
                                for row in table:
                                    content += " | ".join([cell or "" for cell in row]) + "\n"
                                content += "\n"
                    
                    if content.strip():
                        return content
                except Exception as e:
                    logger.warning(f"pdfplumber extraction failed: {e}")
            
            # Fallback to PyPDF2
            if PYPDF2_AVAILABLE:
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        
                        if pdf_reader.is_encrypted:
                            # Try common passwords
                            passwords = options.get('passwords', [''])
                            for password in passwords:
                                try:
                                    pdf_reader.decrypt(password)
                                    break
                                except:
                                    continue
                            else:
                                raise Exception("PDF is password-protected and no valid password provided")
                        
                        for page in pdf_reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                content += page_text + "\n"
                    
                    if content.strip():
                        return content
                except Exception as e:
                    logger.warning(f"PyPDF2 extraction failed: {e}")
            
            # Try OCR if no text found
            if TESSERACT_AVAILABLE and PYMUPDF_AVAILABLE:
                try:
                    doc = fitz.open(file_path)
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        
                        # OCR the image
                        image = Image.open(io.BytesIO(img_data))
                        ocr_text = pytesseract.image_to_string(image, lang='eng')
                        if ocr_text:
                            content += ocr_text + "\n"
                    
                    doc.close()
                    return content
                except Exception as e:
                    logger.warning(f"OCR extraction failed: {e}")
            
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            raise
        
        return content
    
    async def _process_document_advanced(self, file_path: str, options: Dict[str, Any]) -> str:
        """Advanced document processing"""
        content = ""
        
        try:
            if DOCX_AVAILABLE and file_path.endswith('.docx'):
                doc = Document(file_path)
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content += paragraph.text + "\n"
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            content += " | ".join(row_text) + "\n"
                    content += "\n"
            
            elif file_path.endswith('.doc'):
                # For .doc files, we might need to use python-docx2txt or similar
                # For now, return a message
                content = "DOC file processing requires additional libraries. Please convert to DOCX format."
        
        except Exception as e:
            logger.error(f"Document processing error: {e}")
            raise
        
        return content
    
    async def _process_spreadsheet_advanced(self, file_path: str, options: Dict[str, Any]) -> str:
        """Advanced spreadsheet processing"""
        content = ""
        
        try:
            if PANDAS_AVAILABLE:
                # Read all sheets
                excel_file = pd.ExcelFile(file_path)
                
                for sheet_name in excel_file.sheet_names:
                    content += f"\n=== Sheet: {sheet_name} ===\n"
                    
                    # Read sheet data
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Convert to text representation
                    content += df.to_string(index=False) + "\n"
            
            elif OPENPYXL_AVAILABLE:
                workbook = openpyxl.load_workbook(file_path, read_only=True)
                
                for sheet_name in workbook.sheetnames:
                    content += f"\n=== Sheet: {sheet_name} ===\n"
                    sheet = workbook[sheet_name]
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = []
                        for cell in row:
                            if cell is not None:
                                row_text.append(str(cell))
                        if row_text:
                            content += " | ".join(row_text) + "\n"
        
        except Exception as e:
            logger.error(f"Spreadsheet processing error: {e}")
            raise
        
        return content
    
    async def _process_image_advanced(self, file_path: str, options: Dict[str, Any]) -> str:
        """Advanced image processing with OCR"""
        content = ""
        
        try:
            if TESSERACT_AVAILABLE:
                # Get OCR language from options
                lang = options.get('ocr_language', 'eng')
                
                # Process image with OCR
                image = Image.open(file_path)
                
                # Preprocess image if needed
                if options.get('preprocess_image', False):
                    image = self._preprocess_image_for_ocr(image)
                
                # Extract text
                content = pytesseract.image_to_string(image, lang=lang)
                
                # Get confidence scores
                try:
                    data = pytesseract.image_to_data(image, lang=lang, output_type=pytesseract.Output.DICT)
                    confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                    if confidences:
                        avg_confidence = sum(confidences) / len(confidences)
                        logger.info(f"OCR average confidence: {avg_confidence}%")
                except:
                    pass
        
        except Exception as e:
            logger.error(f"Image processing error: {e}")
            raise
        
        return content
    
    async def _process_text_advanced(self, file_path: str, options: Dict[str, Any]) -> str:
        """Advanced text file processing"""
        content = ""
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        break
                except UnicodeDecodeError:
                    continue
            
            if not content:
                # Fallback to binary read
                with open(file_path, 'rb') as file:
                    content = file.read().decode('utf-8', errors='ignore')
        
        except Exception as e:
            logger.error(f"Text processing error: {e}")
            raise
        
        return content
    
    def _preprocess_image_for_ocr(self, image) -> Any:
        """Preprocess image for better OCR results"""
        try:
            # Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')
            
            # Resize if too small
            width, height = image.size
            if width < 300 or height < 300:
                scale_factor = max(300 / width, 300 / height)
                new_size = (int(width * scale_factor), int(height * scale_factor))
                image = image.resize(new_size)
            
            return image
        
        except Exception as e:
            logger.warning(f"Image preprocessing failed: {e}")
            return image
    
    def get_processing_recommendations(self, file_path: str) -> List[str]:
        """Get processing recommendations for a file"""
        recommendations = []
        
        try:
            validation = self.validate_file(file_path)
            
            if not validation['is_valid']:
                return ["File validation failed - cannot process"]
            
            # Add validation recommendations
            recommendations.extend(validation.get('processing_recommendations', []))
            
            # Add size-based recommendations
            file_size_mb = validation['file_info']['size_mb']
            if file_size_mb > 50:
                recommendations.append("Large file - consider splitting or using background processing")
            
            # Add type-specific recommendations
            file_ext = validation['file_info']['extension']
            if file_ext == '.pdf':
                if validation['file_info'].get('num_pages', 0) > 100:
                    recommendations.append("Large PDF - consider page range processing")
            
            elif file_ext in ['.xls', '.xlsx']:
                if validation['file_info'].get('num_sheets', 0) > 5:
                    recommendations.append("Multiple worksheets - consider processing specific sheets")
            
            elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                recommendations.append("Image file - OCR processing will be used")
        
        except Exception as e:
            logger.error(f"Error getting recommendations: {e}")
            recommendations.append("Error analyzing file - proceed with caution")
        
        return recommendations

# Global advanced file processor instance
advanced_file_processor = None

def get_advanced_file_processor() -> AdvancedFileProcessor:
    """Get or create advanced file processor instance"""
    global advanced_file_processor
    if advanced_file_processor is None:
        advanced_file_processor = AdvancedFileProcessor()
    return advanced_file_processor
